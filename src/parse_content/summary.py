import os
import re
from openai import OpenAI
import dotenv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

dotenv.load_dotenv()


def save_as_docx(markdown_text, save_path):
    doc = Document()

    # 设置全文字体为宋体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "SimSun"
    font.size = Pt(12)
    # 强制设置中文字体
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    # 简单的 Markdown 解析
    lines = markdown_text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 标题解析
        if line.startswith("### "):
            h = doc.add_heading(line[4:], level=3)
            # 标题也可以设置字体
            for run in h.runs:
                run.font.name = "SimSun"
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        elif line.startswith("## "):
            h = doc.add_heading(line[3:], level=2)
            for run in h.runs:
                run.font.name = "SimSun"
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        elif line.startswith("# "):
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading.runs:
                run.font.name = "SimSun"
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        # 列表解析
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, line[2:])
        else:
            p = doc.add_paragraph()
            # 设置首行缩进两个字符 (12pt * 2 = 24pt)
            p.paragraph_format.first_line_indent = Pt(24)
            _add_formatted_text(p, line)

    doc.save(save_path)


def _add_formatted_text(paragraph, text):
    # 处理加粗 **text**
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = "SimSun"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        else:
            run = paragraph.add_run(part)
            run.font.name = "SimSun"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def get_summary(article_content, core_idea, essay_motif):
    client = OpenAI(
        api_key=os.environ.get("DEEPSEEK_API_KEY"), base_url="https://api.deepseek.com"
    )

    prompt = """你是一位擅长哲学思辨与议论文写作的助手。我将给你提供三部分内容：

- `原文`：一篇介绍思想家核心观点的文本。
- `核心思想卡`：对原文核心思想的提炼，包括思辨点、内容分析、思辨结论和金句。
- `母题`：一组议论文写作母题，每个母题包含标题、破题角度、思辨性提问和立意进阶要点。

请你根据这些素材，撰写一篇文章议论文写作素材整理的文章。要求如下：

### 1. 文章结构

- **标题**：请为全文拟定一个切合主题、有思想深度且吸引人的标题，应体现思想家的核心观点或文章主旨。
- **开头段**：概括原文思想家的核心观点，引出其思想的重要性或现实意义。**同时，需要点明这位思想家的核心思辨性所在——即他提出了什么关键问题、挑战了什么流行观点、或揭示了什么深刻的悖论，让读者一开始就把握其思想的张力。**
- **主体部分**：针对每一个母题，依次展开论述。每个母题的论述应包含：
  - **小标题**：使用母题的标题。
  - **破题角度**：基于核心思想卡和原文，阐述该母题的切入点，体现思想的深度。
  - **思辨性提问**：提出1-2个具有思辨性的问题，引导读者思考。
  - **立意进阶**：运用核心思想进行更深入的剖析，提出独到的见解，展示如何将思想家的观点应用于母题。

### 2. 写作风格

- 语言应具有学术性和思辨性，但又不失流畅可读。
- 适当引用核心思想卡中的金句或原文的关键表述，增强说服力。
- 避免简单复述原文，要融入自己的理解和发挥，将思想家的观点与现实问题紧密结合。

### 3. 注意事项

- 确保每个母题的论述都紧扣核心思想，不要偏离。
- 思辨性提问要能激发思考，立意进阶要展现思想的延展性。
- 全文逻辑连贯，过渡自然。

### 4. 文章格式范本

```
# [抽象哲理性标题]：[破折号引出现实指向]

[开头段]
- 从现实问题切入，引出思想家的核心观点
- 点明思想家的核心思辨张力（挑战了什么流行观点、揭示了什么悖论）
- 概括核心观点
- 预示文章走向

### 母题一：[母题标题]

[破题角度：约250-350字]
- 基于核心思想阐述该母题的切入点
- 用具体案例或现象阐释抽象思想

[思辨性提问：约100-150字]
- 提出1-2个具有哲学张力的追问
- **问题本身加粗**
- 每个问题可附带简短阐释

[立意进阶：约250-350字]
- 运用核心思想进行更深入的剖析
- 提出独到见解，展示思想如何照亮现实问题
- 可引用金句

### **母题二：[母题标题]**

[同上结构：破题角度→思辨性提问→立意进阶]

...

```
- 注：
  - 每个母题下不要再增加`破题角度`、`思辨性提问`、`立意进阶`这些小标题，直接输出纯文本
  - 每个母题前要有个序号，比如`母题一`、`母题二`等
  - 母题标题不要带`**`

### 5. 素材整理要求

1. **标题**：前半部分为抽象哲学意象，后半部分用破折号指向现实关怀（如“视角的谦逊：在后真相时代重建对话的可能”）

2. **开头段必须包含**：
   - 从现实问题切入，引出思想家的核心观点
   - 明确指出思想家的核心思辨张力——他挑战了什么、揭示了什么悖论

3. **每个母题的论述必须严格遵循三段式结构**（小标题下依次为：破题角度→思辨性提问→立意进阶），段落之间逻辑递进

4. **思辨性提问**需单独成段，问题本身加粗，可附带简短阐释

5. **立意进阶**需超越原文的简单应用，展现思想的延展性和现实穿透力

6. **语言风格**：学术性与可读性平衡，适当引用金句，有对话感，有现实关怀

7. **全文逻辑**：母题之间应有内在关联，共同支撑全文主题

### 6. 示例文章（实际写作中母题数量以输入为准）

```
# [标题]

[开头段内容]

### 母题一：[母题一标题]

[母题一内容]

### 母题二：[母题二标题]

[母题二内容]

...[其他母题]
```

请开始撰写。
"""

    user_content = f"""
---
原文内容
\"\"\"
{article_content}
\"\"\"\n
---
核心思想卡
\"\"\"
{core_idea}
\"\"\"
---
母题
\"\"\"
{essay_motif}
\"\"\"
---
"""
    response = client.chat.completions.create(
        model="deepseek-reasoner",
        messages=[
            {
                "role": "system",
                "content": prompt,
            },
            {"role": "user", "content": user_content},
        ],
        stream=False,
    )

    return response.choices[0].message.content


if __name__ == "__main__":
    core_idea_dir = "dataset/philosophical_proposition/核心思想卡"
    motif_dir = "dataset/philosophical_proposition/母题"
    articles_dir = "dataset/articals"
    summary_dir = "dataset/summary"
    os.makedirs(summary_dir, exist_ok=True)

    for file_name in os.listdir(core_idea_dir):
        if not file_name.endswith(".md"):
            continue

        file_path = os.path.join(core_idea_dir, file_name)
        with open(file_path, "r", encoding="utf-8") as f:
            full_content = f.read()
        with open(os.path.join(motif_dir, file_name), "r", encoding="utf-8") as f:
            essay_motif = f.read()
        ref_match = re.search(
            r"### 参考资料\s*[:：]\s*\n\s*-\s*(.*?)\((.*?)\)[:：]\s*(.*)", full_content
        )

        article_content = ""
        clean_content = full_content

        if ref_match:
            book = ref_match.group(1).strip()
            author = ref_match.group(2).strip()
            chapter = ref_match.group(3).strip()

            # 去除参考资料部分的内容作为核心思想卡内容
            clean_content = full_content[: ref_match.start()].strip()

            # 读取原文
            article_filename = f"{chapter}.md"
            article_path = os.path.join(articles_dir, article_filename)

            if os.path.exists(article_path):
                with open(article_path, "r", encoding="utf-8") as af:
                    article_content = af.read()
            else:
                print(f"警告：未找到原文文件 {article_path}")
        else:
            print(f"注意：文件 {file_name} 未能解析出参考资料")

        print(f"正在处理: {file_name}...")
        response = get_summary(
            core_idea=clean_content,
            article_content=article_content,
            essay_motif=essay_motif,
        )

        # 添加引用来源
        if ref_match:
            book = ref_match.group(1).strip()
            author = ref_match.group(2).strip()
            chapter = ref_match.group(3).strip()
            ref_footer = f"\n\n---\n**参考资料**：\n- 《{book}》 ({author}) : {chapter}"
            response += ref_footer

        # 保存为 Markdown
        save_file_name = f"{chapter}_{file_name}"
        md_save_path = os.path.join(summary_dir, save_file_name)
        with open(md_save_path, "w", encoding="utf-8") as f:
            f.write(response)

        # 保存为 Word
        docx_save_path = os.path.join(
            summary_dir, save_file_name.replace(".md", ".docx")
        )
        try:
            save_as_docx(response, docx_save_path)
            print(f"已保存: {md_save_path} 和 {docx_save_path}")
        except Exception as e:
            print(f"保存 Word 文件失败: {e}")
